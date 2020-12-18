import string
from flask import Flask, request, render_template, redirect, session, url_for, send_file, flash
from flask_session import Session
from tempfile import mkdtemp
from werkzeug.exceptions import default_exceptions, HTTPException, InternalServerError
from werkzeug.security import check_password_hash, generate_password_hash
from sqlalchemy import create_engine, MetaData, Table, Column, Integer, String, ColumnDefault, ForeignKey, select
from docx import Document

from helpers import apology, login_required

app = Flask(__name__)
engine = create_engine('sqlite:///handout_maker.db')
metadata = MetaData(bind=engine)

app.config["SESSION_FILE_DIR"] = mkdtemp()
app.config["SESSION_PERMANENT"] = False
app.config["SESSION_TYPE"] = "filesystem"
Session(app)


user_table = Table('users', metadata,
                   Column('id', Integer, primary_key=True, autoincrement=True),
                   Column('username', String(30), index=True, nullable=False, unique=True),
                   Column('hash', String, nullable=False),
                   Column('name', String(30), nullable=False),
                   Column('last_name', String(30), nullable=False),
                   Column('handouts', Integer))    

handout_table = Table('handouts', metadata,
                   Column('handout_id', Integer, primary_key=True, autoincrement=True),  
                   Column('user_id', Integer),
                   Column('handout_name', String(30), index=True, nullable=False),     
                   Column('language', String(30), index=True, nullable=False),                   
                   Column('level', Integer, nullable=False),
                   Column('desc_csv', String, nullable=False))

lesson_table = Table('lessons', metadata,
                   Column('lesson_id', Integer, primary_key=True, autoincrement=True),                        
                   Column('handout_id', Integer),                     
                   Column('user_id', Integer),              
                   Column('lesson_num', Integer, index=True, nullable=False),
                   Column('title', String(30), index=True, nullable=False),
                   Column('text', String, nullable=False),
                   Column('exercise', String),
                   Column('text_two', String),
                   Column('morphology', String),
                   Column('syntax', String),
                   Column('tips', String))

vocabulary_table = Table('vocabularies', metadata,
                   Column('vocabulary_id', Integer, primary_key=True, autoincrement=True),                       
                   Column('lesson_id', Integer), 
                   Column('handout_id', Integer), 
                   Column('text', Integer),
                   Column('word', String(45)),                   
                   Column('type', String(20)),            
                   Column('gen', String(1)),
                   Column('translate', String(46)))   

@app.route("/")
def index():
    return redirect("/myhandout")

@app.route("/login", methods=['GET', 'POST'])
def login():
    session.clear()

    if request.method == "POST":
        
        if not request.form.get("username"):
            return apology("must provide username", 403)

        elif not request.form.get("password"):
            return apology("must provide password", 403)

        rows = engine.execute("SELECT * FROM users WHERE username = :username",
                          username=request.form.get("username")).fetchall()

        if len(rows) != 1 or not check_password_hash(rows[0]["hash"], request.form.get("password")):
            return apology("invalid username and/or password", 403)

        session["user_id"] = rows[0]["id"]

        return redirect("/myhandout")
    
    return render_template("login.html")

@app.route("/register", methods=['GET', 'POST'])
def register():
    session.clear()

    # User reached route via POST (as by submitting a form via POST)
    if request.method == "POST":
        name = request.form.get("name")
        last_name = request.form.get("last_name")
        user = request.form.get("username")
        password = request.form.get("password")
        passwordagain = request.form.get("passwordagain")
        handouts = 0

        # Ensure password not equals password
        if password != passwordagain:
            return apology("password not equals", 404)

        passhash = generate_password_hash(password)

        # Ensure username was submitted
        if not user:
            return apology("must provide username", 403)

        # Ensure password was submitted
        elif not password:
            return apology("must provide password", 403)

        else:
            with engine.connect() as conn:
                db = conn.execute("INSERT INTO users(username, hash, name, last_name, handouts) VALUES(:username, :password, :name, :last_name, :handouts)", 
                        username=user, password=passhash, name=name, last_name=last_name, handouts=handouts)

                # Query database for username
                rows = conn.execute("SELECT * FROM users WHERE username = :username",
                                           username=request.form.get("username")).fetchall()
                       
                if len(rows) != 1 or not check_password_hash(rows[0]["hash"], request.form.get("password")):
                    return apology("invalid username and/or password", 403)

            session["user_id"] = rows[0]["id"]
            return redirect("/myhandout")

    return render_template("/register.html")

@app.route("/logout")
def logout():
    # Forget any user_id
    session.clear()

    # Redirect user to login form
    return redirect("/myhandout")

@app.route("/newhandout", methods=['GET', 'POST'])
@login_required
def create_handout():    
    if request.method == "POST":
        handout_name = request.form.get("handout_name")             
        level = request.form.get("level")   
        language = request.form.get("language")        
        desc_csv = request.form.get("desc_csv")        
        with engine.connect() as conn:
            conn.execute("INSERT INTO handouts(user_id, handout_name, language, level, desc_csv) VALUES(?, ?, ?, ?, ?)",
                                                (session['user_id'], handout_name, language, level, desc_csv))
            select_handouts = conn.execute("SELECT handouts FROM users WHERE id = ?", (session['user_id'],)).fetchall()
            handouts_qtd = int(select_handouts[0][0])
            handouts_qtd += 1
            update_handouts = conn.execute("UPDATE users SET handouts = ? WHERE id = ?", (handouts_qtd, session['user_id'],))
           
        return redirect("/myhandout")
         
    return render_template("new_handout.html")  
    
@app.route("/myhandout")
@login_required
def myhandout():
    # DATA OF INDEX
    ids = []
    names = []
    languages = []
    levels = []
    desc_csv = []
    with engine.connect() as conn:
        db_handout = conn.execute("SELECT * FROM handouts WHERE user_id = ?", (session['user_id'])).fetchall()
        count_handout = conn.execute("SELECT COUNT(*) as count FROM handouts WHERE user_id = ?", (session['user_id'])).fetchall()
        count = count_handout[0]['count']
        for i in range(count):             
            data_id = db_handout[i]['handout_id']
            data_name = db_handout[i]['handout_name']
            data_language = db_handout[i]['language']
            data_level = db_handout[i]['level']
            data_desc = db_handout[i]['desc_csv']
            ids.append(data_id)
            names.append(data_name)
            languages.append(data_language)
            levels.append(data_level)
            desc_csv.append(data_desc)
        
    return render_template("myhandout.html", count=count, name=names, language=languages, level=levels, desc=desc_csv, handout_id=ids)

@app.route("/handout/<int:handout_id>/")
@login_required
def handout(handout_id):
    with engine.connect() as conn:        
        db_handout = conn.execute("SELECT * FROM lessons WHERE handout_id = ? AND user_id = ? ORDER BY lesson_num DESC", (handout_id, session['user_id'])).fetchall()
        db_handoutcount = conn.execute("SELECT COUNT(*) as count FROM lessons WHERE handout_id = ? AND user_id = ?", (handout_id, session['user_id'])).fetchall()
        db_handoutname = conn.execute("SELECT handout_name FROM handouts WHERE handout_id = ?", (handout_id)).fetchall()
        size = db_handoutcount[0]['count']          
        title_pag = db_handoutname[0]['handout_name']  
        try:                        
            get_title = []
            get_number = []
            get_handout = []
            get_id = []
            for i in range(0, size):
                get_title.append(db_handout[i]['title'])
                get_number.append(db_handout[i]['lesson_num'])                
                get_id.append(db_handout[i]['lesson_id'])
                get_handout.append(db_handoutname[0]['handout_name'])
        except: 
            return render_template('handout.html', size=size, title=get_title, number=get_number, handout_id=handout_id, lesson_id=get_id, handout_name=get_handout, title_pag=title_pag)
    return render_template('handout.html', size=size, title=get_title, number=get_number, handout_id=handout_id, lesson_id=get_id, handout_name=get_handout, title_pag=title_pag)

@app.route("/edit/handout/<int:handout_id>/", methods=['GET', 'POST'])
@login_required
def edit_handout(handout_id): 
    if request.method == "POST":
        handout_name = request.form.get("handout_name")             
        level = request.form.get("level")   
        language = request.form.get("language")        
        desc_csv = request.form.get("desc_csv")        
        with engine.connect() as conn:
            update = conn.execute("UPDATE handouts SET handout_name = ?, level = ?, language = ?, desc_csv = ? WHERE handout_id = ? AND user_id = ?", 
                                  (handout_name, level, language, desc_csv, handout_id, session['user_id']))
                                                
        return redirect("/myhandout")
    with engine.connect() as conn: 
        db_handout = conn.execute("SELECT * FROM handouts WHERE handout_id = ? AND user_id = ?", (handout_id, session['user_id'])).fetchall()
        handout_name = db_handout[0]['handout_name']
        handout_level = db_handout[0]['level']
        handout_language = db_handout[0]['language']
        handout_desc = db_handout[0]['desc_csv']
    
    return render_template('edit_handout.html', handout_id=handout_id, handout_name=handout_name, handout_level=handout_level, handout_language=handout_language, handout_desc=handout_desc)

@app.route("/newlesson", methods=['GET', 'POST'])
@login_required
def create_lesson():    
    with engine.connect() as conn:
        db_handout = conn.execute("SELECT handout_name FROM handouts WHERE user_id = ?", (session['user_id'])).fetchall()
        db_handoutcount = conn.execute("SELECT COUNT(*) as count FROM handouts WHERE user_id = ?", (session['user_id'])).fetchall()
        size = db_handoutcount[0]['count']         
        get_handouts = []

        for i in range(0, size):    
            get_handouts.append(db_handout[i]['handout_name'] )

    if request.method == "POST":
        handout_name = request.form.get("handout_newlesson")
        title_lesson = request.form.get("title_lesson")
        num_lesson = request.form.get("num_lesson")
        main_text = request.form.get("main_text")
        exercises = request.form.get("exercises")
        text_two = request.form.get("text_two")
        morphology = request.form.get("morphology")
        syntax = request.form.get("syntax")
        tips = request.form.get("tips")
        handout_id = 0
        with engine.connect() as conn:
            db_handoutid = conn.execute("SELECT handout_id FROM handouts WHERE handout_name = ? AND user_id = ?", (handout_name, session['user_id'],)).fetchall()

            handout_id = db_handoutid[0]['handout_id']
            db_lesson = conn.execute("INSERT INTO lessons(handout_id, title, lesson_num, text, exercise, text_two, morphology, syntax, tips, user_id)VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?)", (handout_id,  title_lesson, num_lesson, main_text, exercises, text_two, morphology, syntax, tips, session['user_id'],))
        # VOCABULARY
        vocabulary_list = []
        with engine.connect() as conn:
            vocabulary = conn.execute("SELECT word FROM vocabularies JOIN handouts WHERE handout_name LIKE ? AND user_id = ? AND vocabularies.handout_id = ?", (handout_name, session['user_id'], handout_id)).fetchall()          
            vocabulary_count = conn.execute("SELECT COUNT(word) as count FROM vocabularies JOIN handouts WHERE handout_name LIKE ? AND user_id = ? AND vocabularies.handout_id = ?", (handout_name, session['user_id'], handout_id)).fetchall()
            v_size = vocabulary_count[0]['count']
            for v in range(v_size): 
                for linha in vocabulary[v]: 
                    linha = linha.rstrip()
                    linha = linha.translate(linha.maketrans('', '', string.punctuation))
                    linha = linha.lower()
                    palavras = linha.split()
                    for palavra in palavras:        
                        if palavra not in vocabulary_list:
                            vocabulary_list.append(vocabulary[v]['word'])

            texts = conn.execute("SELECT text FROM lessons WHERE user_id = ?", (session['user_id'])).fetchall()    
            text_info = conn.execute("SELECT handout_id, lesson_id FROM lessons WHERE user_id = ?", (session['user_id'])).fetchall()    
            text_count = conn.execute("SELECT COUNT(text) as count FROM lessons WHERE user_id = ?", (session['user_id'])).fetchall()
            size = text_count[0]['count']  
            for i in range(size):  
                for line in texts[i]: 
                    line = line.rstrip()
                    line = line.translate(line.maketrans('', '', string.punctuation))
                    line = line.lower()
                    words = line.split()
                    for word in words:      
                        if word not in vocabulary_list:                        
                            vocabulary_list.append(word)
                            db_vocabulary = conn.execute("INSERT INTO vocabularies(word, handout_id, lesson_id, text) VALUES(?, ?, ?, ?)", 
                                                        (word, text_info[i]['handout_id'], text_info[i]['lesson_id'], handout_name))
        return redirect("/myhandout")
    return render_template("new_lesson.html", size=size, handout_op=get_handouts)

@app.route("/edit/lesson/<int:lesson_id>/", methods=['GET', 'POST'])
@login_required
def edit_lesson(lesson_id):
    if request.method == "POST":
        handout_name = request.form.get("handout_newlesson")
        title_lesson = request.form.get("title_lesson")
        num_lesson = request.form.get("num_lesson")
        main_text = request.form.get("main_text")
        exercises = request.form.get("exercises")
        text_two = request.form.get("text_two")
        morphology = request.form.get("morphology")
        syntax = request.form.get("syntax")
        tips = request.form.get("tips")

        with engine.connect() as conn:
            del_vocabulary = conn.execute('''DELETE FROM vocabularies WHERE word IN 
                                            (SELECT word FROM vocabularies JOIN handouts 
                                             WHERE handout_name LIKE ? AND user_id = ? AND lesson_id = ?)''',
                                            (handout_name, session['user_id'], lesson_id))

            update_lesson = conn.execute('''UPDATE lessons SET title = ?, lesson_num = ?, 
                                            text = ?, exercise = ?, text_two = ?, morphology = ?,
                                            syntax = ?, tips = ?
                                            WHERE user_id = ? AND lesson_id = ?''',
                                            title_lesson, num_lesson, main_text, exercises, text_two, 
                                            morphology, syntax, tips, 
                                            session['user_id'], lesson_id)
        # VOCABULARY
        vocabulary_list = []
        with engine.connect() as conn:
            vocabulary = conn.execute("SELECT word FROM vocabularies JOIN handouts WHERE handout_name LIKE ? AND user_id = ?", (handout_name, session['user_id'])).fetchall()          
            vocabulary_count = conn.execute("SELECT COUNT(word) as count FROM vocabularies JOIN handouts WHERE handout_name LIKE ? AND user_id = ?", (handout_name, session['user_id'])).fetchall()
            v_size = vocabulary_count[0]['count']
            for v in range(v_size): 
                for linha in vocabulary[v]: 
                    linha = linha.rstrip()
                    linha = linha.translate(linha.maketrans('', '', string.punctuation))
                    linha = linha.lower()
                    palavras = linha.split()
                    for palavra in palavras:        
                        if palavra not in vocabulary_list:
                            vocabulary_list.append(vocabulary[v]['word'])

            texts = conn.execute("SELECT text FROM lessons WHERE user_id = ?", (session['user_id'])).fetchall()    
            text_info = conn.execute("SELECT handout_id, lesson_id FROM lessons WHERE user_id = ?", (session['user_id'])).fetchall()    
            text_count = conn.execute("SELECT COUNT(text) as count FROM lessons WHERE user_id = ?", (session['user_id'])).fetchall()
            size = text_count[0]['count']  
            for i in range(size):  
                for line in texts[i]: 
                    line = line.rstrip()
                    line = line.translate(line.maketrans('', '', string.punctuation))
                    line = line.lower()
                    words = line.split()
                    for word in words:      
                        if word not in vocabulary_list:                        
                            vocabulary_list.append(word)
                            db_vocabulary = conn.execute("INSERT INTO vocabularies(word, handout_id, lesson_id, text) VALUES(?, ?, ?, ?)", 
                                                        (word, text_info[i]['handout_id'], text_info[i]['lesson_id'], handout_name))
        return redirect("/myhandout")
    with engine.connect() as conn:
        db_handout = conn.execute("SELECT handout_name FROM handouts WHERE user_id = ?", (session['user_id'])).fetchall()
        db_handoutcount = conn.execute("SELECT COUNT(*) as count FROM handouts WHERE user_id = ?", (session['user_id'])).fetchall()
        size = db_handoutcount[0]['count']         
        get_handouts = []

        for i in range(0, size):    
            get_handouts.append(db_handout[i]['handout_name'] )
    return render_template('edit_lesson.html', lesson_id=lesson_id, size=size, handout_op=get_handouts)

@app.route("/delete/lesson/<int:handout_id>/<int:lesson_id>/")
@login_required
def delete_lesson(handout_id, lesson_id):
    with engine.connect() as conn:
        del_lesson = conn.execute('''DELETE FROM lessons WHERE user_id = ? AND lesson_id = ?''',
                                    session['user_id'], lesson_id)
            
        del_vocabulary = conn.execute('''DELETE FROM vocabularies WHERE word IN 
                                            (SELECT word FROM vocabularies JOIN handouts
                                             WHERE vocabularies.handout_id = ? AND lesson_id = ?)''',
                                            (handout_id, lesson_id))
    return redirect('/myhandout')

@app.route("/vocabulary", methods=['GET', 'POST'])
@login_required
def myvocabulary():
    if request.method == "POST":
        handout_name = request.form.get("handout_name")

        with engine.connect() as conn:
            db_handoutid = conn.execute("SELECT handout_id FROM handouts WHERE handout_name = ? AND user_id = ?", (handout_name, session['user_id'],)).fetchall()
            handout_id = db_handoutid[0]['handout_id'] 
            return redirect(url_for('vocabulary', handout_id=handout_id))
                        
    with engine.connect() as conn:
        db_handout = conn.execute("SELECT handout_name FROM handouts WHERE user_id = ?", (session['user_id'])).fetchall()
        db_handoutcount = conn.execute("SELECT COUNT(*) as count FROM handouts WHERE user_id = ?", (session['user_id'])).fetchall()
        size = db_handoutcount[0]['count']         
        get_handouts = []

        for i in range(0, size):    
            get_handouts.append(db_handout[i]['handout_name'] )

    return render_template('vocabulary.html', size=size, handout_name=get_handouts)

@app.route("/vocabulary/<int:handout_id>/")
@login_required
def vocabulary(handout_id):
    with engine.connect() as conn:        
        db_vocabulary = conn.execute("SELECT * FROM vocabularies JOIN handouts WHERE vocabularies.handout_id = ? AND user_id = ?", (handout_id, session['user_id'])).fetchall()
        db_vocabularycount = conn.execute("SELECT COUNT(*) as count FROM vocabularies JOIN handouts WHERE vocabularies.handout_id = ? AND user_id = ?", (handout_id, session['user_id'])).fetchall()

        size = db_vocabularycount[0]['count']           
        try:                        
            get_word = []
            get_translate = []
            get_id = []
            for i in range(0, size):
                get_word.append(db_vocabulary[i]['word'])
                get_translate.append(db_vocabulary[i]['translate'])                
                get_id.append(db_vocabulary[i]['vocabulary_id'])
        except: 
            return render_template('vocabulary.html')
    return render_template('vocabulary_id.html', size=size, word=get_word, translate=get_translate, vocabulary_id=get_id)

@app.route("/vocabulary/<int:handout_id>/<int:vocabulary_id>", methods=['GET', 'POST'])
@login_required
def edit_vocabulary(handout_id, vocabulary_id):
    with engine.connect() as conn:        
        db_vocabulary = conn.execute('''SELECT * FROM vocabularies JOIN handouts 
                                        WHERE vocabularies.handout_id = ? AND user_id = ? AND vocabulary_id = ?''', 
                                        (handout_id, session['user_id'], vocabulary_id)).fetchall()
        title = db_vocabulary[0]['text']
        word = db_vocabulary[0]['word']

    if request.method == "POST": 
        translate = request.form.get("translate") 
        type_word = request.form.get("type") 
        gen_word = request.form.get("gen") 
        with engine.connect() as conn: 
            db_vocabulary = conn.execute('''UPDATE vocabularies 
                                            SET translate = ?, type = ?, gen = ? 
                                            WHERE vocabularies.handout_id = ? AND vocabulary_id = ?''', 
                                            (translate, type_word, gen_word,
                                            handout_id, vocabulary_id))
        return redirect("/vocabulary")
    return render_template('edit_vocabulary.html', word=word, title=title)

@app.route("/vocabulary/delete/<int:vocabulary_id>")
@login_required
def delete_vocabulary(vocabulary_id):
    with engine.connect() as conn:
        db_vocabulary = conn.execute('''DELETE FROM vocabularies WHERE word IN 
                                        (SELECT word FROM vocabularies JOIN handouts 
                                        WHERE user_id = ? AND vocabulary_id = ?)''', 
                                        (session['user_id'], vocabulary_id))
    return redirect("/vocabulary")

@app.route("/readability", methods=["GET", "POST"])
def readability():
    if request.method == "POST":
        text = request.form.get("text")
        if len(text) == 0:
            grade = "Input a text"
        else:     
            # Lenght of Letters
            count_punctuation = 0
            for p in range(len(text)):
                if (text[p] in string.punctuation):
                    count_punctuation += 1

            letters = len(text) - text.count(" ") - count_punctuation

            # Lenght of sentences
            sentences = 0
            for p in range(len(text)):
                if (text[p] == ".") or (text[p] == "!") or (text[p] == "?"):
                    sentences += 1

            # Lenght of words
            words = len(text.split())

            # Index Coleman-Liau
            L = letters / words * 100
            S = sentences / words * 100
            CLI = (0.0588 * L) - (0.296 * S) - 15.8
            CLI = round(CLI)
            
            # Output 
            if CLI <= 1:
                grade = "Before Grade 1"
            elif CLI >= 16:
                grade = "Grade 16+"
            else:
                grade = f"Grade {CLI}"
        return render_template('readability.html', text=text, grade=grade)

    return render_template('readability.html')

@app.route("/pdf/generate/<int:handout_id>", methods=['GET', 'POST'])
@login_required
def generate_pdf(handout_id):
    try:
        with engine.connect() as conn:
        
            db_count =conn.execute('''SELECT COUNT(*) as count 
                                    FROM handouts JOIN lessons
                                    WHERE handouts.user_id = ? AND handouts.handout_id = ? AND lessons.handout_id = ?''',
                                    session['user_id'], handout_id, handout_id).fetchall()
            dados = conn.execute('''SELECT handout_name, language, lesson_num, title, text, exercise, morphology, syntax, tips   
                                    FROM handouts JOIN lessons 
                                    WHERE handouts.user_id = ? AND handouts.handout_id = ? AND lessons.handout_id = ?''',
                                    session['user_id'], handout_id, handout_id).fetchall()
            handout_name=dados[0]['handout_name']
            language = dados[0]['language']
            count = db_count[0]['count']
            print(f'count: {count}')
            get_lesson = []
            get_title = []
            get_text = []
            get_exercise = []
            get_morphology = []
            get_syntax = []
            get_tips = []        
            get_countvoc = []
            get_word = [[]]
            get_type = [[]]
            get_gen = [[]]
            get_translate = [[]]

            # Open Document 
            document = Document()
            # First page
            doc_title = document.add_heading(handout_name, 0)
            doc_language = document.add_paragraph(language)
            document.add_page_break()

            for i in range(count):
                get_lesson.append(dados[i]['lesson_num'])
                get_title.append(dados[i]['title'])
                get_text.append(dados[i]['text'])
                get_exercise.append(dados[i]['exercise'])
                get_morphology.append(dados[i]['morphology'])
                get_syntax.append(dados[i]['syntax'])
                get_tips.append(dados[i]['tips'])
                
                print(f'i:{i}')

                vocabulary_db = conn.execute('''SELECT word, type, gen, translate 
                                                FROM vocabularies JOIN handouts 
                                                WHERE handouts.user_id = ? AND handouts.handout_id = ? AND lesson_id = ?
                                                ORDER BY type''',
                                                session['user_id'], handout_id, i+1).fetchall()
                vocabulary_count = conn.execute('''SELECT COUNT(*) as count_voc
                                                    FROM vocabularies JOIN handouts 
                                                    WHERE handouts.user_id = ? AND handouts.handout_id = ? AND lesson_id = ?
                                                    ORDER BY type''',
                                                    session['user_id'], handout_id, i+1).fetchall()
                get_countvoc.append(vocabulary_count[0]['count_voc'])   
                count_voc = vocabulary_count[0]['count_voc']
                print(f"get_countvoc: {get_word}")
                print(f"vocabulary: {count_voc}")    

                for v in range(count_voc):

                    get_word[i].append(vocabulary_db[v]['word'])
                    get_type[i].append(vocabulary_db[v]['type'])
                    get_gen[i].append(vocabulary_db[v]['gen'])
                    get_translate[i].append(vocabulary_db[v]['translate'])
                    print(f'v: {v}')
                    word = vocabulary_db[v]['word']
                    print(f'word: {word}')
                
                get_word.append([])
                get_type.append([])
                get_gen.append([])
                get_translate.append([])

        for d in range(count):
            # Page of Lesson
            document.add_heading(f"Lesson {get_lesson[d]}", level=1)        
            document.add_heading(get_title[d], level=2)        
            document.add_paragraph(get_text[d])

            table = document.add_table(rows=1, cols=4)      
            voc_cells = table.rows[0].cells  
            voc_cells[0].text = 'Word'
            voc_cells[1].text = 'Type'
            voc_cells[2].text = 'Gen'            
            voc_cells[3].text = 'Translate'
            for voc_docs in range(get_countvoc[d]):
                row_cells = table.add_row().cells
                try: 
                    row_cells[0].text = get_word[d][voc_docs]
                except:
                    row_cells[0].text = "None"
                try:                 
                    row_cells[1].text = get_type[d][voc_docs]
                except:
                    row_cells[1].text = "None"
                try:                
                    row_cells[2].text = get_gen[d][voc_docs]
                except:
                    row_cells[2].text = "None"
                try:                 
                    row_cells[3].text = get_translate[d][voc_docs]
                except:
                    row_cells[3].text = "None"

                
            document.add_heading('Exercises', level=2)
            document.add_paragraph(get_exercise[d])    
            
            document.add_page_break()

            # Page of Grammar 
            if len(get_morphology[d]) != 0:
                document.add_heading('Morphology', level=2)
                document.add_paragraph(get_morphology[d])  
            if len(get_syntax[d]) != 0:
                document.add_heading('Syntax', level=2)
                document.add_paragraph(get_syntax[d])   
            if len(get_tips[d]) != 0:
                document.add_heading('Tips', level=2)
                document.add_paragraph(get_tips[d])  
            
            document.add_page_break()

        document.save(f'{handout_name}.docx')
        arquivo = f'{handout_name}.docx'
        
        if request.method == 'POST':
            return send_file(arquivo,as_attachment=True)
    except:
        flash('Your handout is empty. Take new lessons!')
        return redirect(f'/handout/{handout_id}')
    return redirect ('/myhandout')

if __name__ == "__main__":
    metadata.create_all()
    app.run(port=8080, host='127.0.0.1', debug=True, threaded=True)
